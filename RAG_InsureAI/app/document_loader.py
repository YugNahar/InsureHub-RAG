"""
document_loader.py — Unified document loading for InsureHub RAG.

Supported sources
-----------------
  PDFs          : PyMuPDF (fitz) with multi-page text extraction
  DOCX          : python-docx
  XLSX / CSV    : pandas
  Images        : pytesseract OCR (optional)
  Audio/Video   : OpenAI Whisper (optional, for local transcription)
  YouTube URLs  : YouTubeTranscriptApi >= 0.6 + 800-word overlapping chunks
  Web URLs      : requests + BeautifulSoup HTML → text
                  Playwright fallback for JS-rendered (SPA) pages

Public API
----------
  load_document(file_path, filename)  → list[Document]
  load_url(url)                       → list[Document]
  load_url_advanced(url, **kw)        → list[Document]
  extract_urls(text)                  → list[str]
  is_youtube_url(url)                 → bool
  _load_youtube(url)                  → list[Document]  (also used by eval_api)
  _get_whisper_model()                → whisper.Model | None

Constants
---------
  ALLOWED_EXTENSIONS  : set of lowercase extensions the API accepts
  MAX_FILE_SIZE_BYTES : 50 MB hard limit
  FileValidationError : exception raised for invalid files

Webpage cleaning improvements (v2)
-----------------------------------
  1. Playwright fallback  — fetches JS-rendered SPA pages that return empty
                            HTML to plain requests.
  2. Encoding detection   — uses chardet to decode non-UTF-8 pages correctly
                            before handing the markup to BeautifulSoup.
  3. Nav-link noise filter— drops lines shorter than 4 words that are almost
                            always breadcrumbs / menu links / cookie banners.
  4. Friendly HTTP errors — 403/429/5xx give an actionable message instead of
                            a bare RuntimeError.
  5. Duplicate-line dedup — consecutive identical lines (e.g. repeated headers
                            rendered by JS frameworks) are collapsed to one.
"""
from __future__ import annotations

import logging
import os
import re
import tempfile
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

# ── LangChain document type ────────────────────────────────────────────────────
try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.schema import Document  # type: ignore

# ── Constants ─────────────────────────────────────────────────────────────────
ALLOWED_EXTENSIONS: set[str] = {
    ".pdf", ".docx", ".doc", ".txt", ".csv", ".xlsx", ".xls",
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp",
    ".mp3", ".mp4", ".wav", ".m4a", ".ogg", ".webm",
}

MAX_FILE_SIZE_BYTES: int = 50 * 1024 * 1024  # 50 MB
_SAFE_TMP_ROOT = os.path.expanduser("~/.insurehub_tmp")
os.makedirs(_SAFE_TMP_ROOT, exist_ok=True)

# Minimum meaningful line length for webpage noise filtering (words)
_MIN_LINE_WORDS: int = 4

# If requests returns fewer than this many chars of body text, try Playwright
_SPARSE_CONTENT_THRESHOLD: int = 300


class FileValidationError(ValueError):
    """Raised when an uploaded file fails validation."""


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def extract_urls(text: str) -> list[str]:
    """Extract all HTTP/HTTPS URLs from a string."""
    pattern = r"https?://[^\s\"'<>)\]}\|\\]+"
    return re.findall(pattern, text)


def is_youtube_url(url: str) -> bool:
    """Return True if *url* is a YouTube video link."""
    try:
        parsed = urlparse(url)
        host = parsed.netloc.lower().lstrip("www.")
        if host == "youtube.com":
            return "watch" in parsed.path or "embed" in parsed.path or "shorts" in parsed.path
        if host == "youtu.be":
            return bool(parsed.path.strip("/"))
        return False
    except Exception:
        return False


def _extract_video_id(url: str) -> Optional[str]:
    """Extract the YouTube video ID from any supported URL format."""
    patterns = [
        r"(?:v=|/v/|youtu\.be/|/embed/|/shorts/)([A-Za-z0-9_-]{11})",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def _detect_encoding(raw_bytes: bytes) -> str:
    """
    Detect the character encoding of raw HTTP response bytes.
    Falls back to utf-8 if chardet is not installed or detection fails.
    """
    try:
        import chardet
        result = chardet.detect(raw_bytes[:10_000])
        encoding = result.get("encoding") or "utf-8"
        # chardet sometimes returns 'ascii' for utf-8 content — promote it
        return "utf-8" if encoding.lower() == "ascii" else encoding
    except ImportError:
        return "utf-8"


def _clean_webpage_text(raw_text: str) -> str:
    """
    Post-process plain text extracted from a web page:
      1. Drop lines shorter than _MIN_LINE_WORDS words  (nav / breadcrumb noise)
      2. Remove consecutive duplicate lines              (JS-framework repeated headers)
      3. Collapse runs of 3+ blank lines into 2          (already done by caller, kept for safety)
      4. Strip leading/trailing whitespace
    """
    lines = raw_text.splitlines()
    cleaned: list[str] = []
    prev_stripped = None
    for line in lines:
        stripped = line.strip()
        # Keep blank lines (they preserve paragraph structure) but collapse dups
        if not stripped:
            if prev_stripped != "":
                cleaned.append("")
            prev_stripped = ""
            continue
        # Drop very short lines — almost always navigation / cookie / breadcrumb
        if len(stripped.split()) < _MIN_LINE_WORDS:
            continue
        # Drop consecutive duplicate lines
        if stripped == prev_stripped:
            continue
        cleaned.append(stripped)
        prev_stripped = stripped

    text = "\n".join(cleaned)
    # Final whitespace collapse
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


# ══════════════════════════════════════════════════════════════════════════════
# WHISPER (optional audio transcription)
# ══════════════════════════════════════════════════════════════════════════════
_whisper_model = None
_whisper_loaded = False


def _get_whisper_model():
    """Lazy-load the Whisper 'base' model. Returns None if whisper is not installed."""
    global _whisper_model, _whisper_loaded
    if _whisper_loaded:
        return _whisper_model
    _whisper_loaded = True
    try:
        import whisper
        _whisper_model = whisper.load_model("base")
        logger.info("[Whisper] Model 'base' loaded.")
    except Exception as exc:
        logger.warning("[Whisper] Not available: %s", exc)
        _whisper_model = None
    return _whisper_model


# ══════════════════════════════════════════════════════════════════════════════
# YOUTUBE LOADER
# ══════════════════════════════════════════════════════════════════════════════

def _clean_youtube_transcript(text: str) -> str:
    """
    Clean YouTube auto-generated transcript text.
    Removes filler words and deduplicates sentences that appear twice in a row
    (a common YouTube auto-caption artifact where each spoken sentence is
    repeated in the next caption segment).
    """
    # Remove filler words/phrases (case-insensitive, whole word only)
    filler_patterns = [
        r"\buh\b", r"\bum\b", r"\buhh\b", r"\bumm\b",
        r"\byou know\b", r"\bi mean\b", r"\blike\b(?=\s+\w)",
        r"\bso\b(?=\s*,|\s*$)", r"\bright\b(?=\s*,|\s*$)",
        r"\bokay so\b", r"\balright so\b",
    ]
    for pat in filler_patterns:
        text = re.sub(pat, "", text, flags=re.IGNORECASE)

    # Collapse multiple spaces
    text = re.sub(r"  +", " ", text).strip()

    # Deduplicate consecutive repeated sentences/phrases.
    # YouTube auto-captions often repeat each clause in the next segment:
    # "... basis of mutual solidarity  basis of mutual solidarity ..."
    # We detect this by splitting into sentences and removing back-to-back dups.
    sentences = re.split(r"(?<=[.!?])\s+", text)
    deduped: list[str] = []
    prev = None
    for s in sentences:
        s_norm = s.strip().lower()
        if s_norm and s_norm != prev:
            deduped.append(s.strip())
            prev = s_norm

    # Also catch partial overlapping repetition at phrase level (split on commas/conjunctions)
    rejoined = " ".join(deduped)
    # Remove exact substring repetitions at word-ngram level (window of 8 words)
    words = rejoined.split()
    result_words: list[str] = []
    i = 0
    while i < len(words):
        # Look for a repeated window of 6+ words starting from position i
        found_dup = False
        for wlen in range(min(12, len(words) - i), 5, -1):
            window = words[i:i + wlen]
            # Check if this window appears just after current position
            next_start = i + wlen
            if next_start + wlen <= len(words):
                next_window = words[next_start:next_start + wlen]
                if window == next_window:
                    result_words.extend(window)
                    i += wlen * 2  # skip both copies
                    found_dup = True
                    break
        if not found_dup:
            result_words.append(words[i])
            i += 1

    cleaned = " ".join(result_words)
    # Final whitespace cleanup
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _get_video_title(url: str, video_id: str) -> str:
    """Fetch YouTube video title via the public oEmbed API (no auth required)."""
    try:
        import requests as _req
        oembed = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
        resp = _req.get(oembed, timeout=6)
        if resp.ok:
            title = resp.json().get("title", "").strip()
            if title:
                logger.info("[YouTube] Video title from oEmbed: %r", title)
                return title
    except Exception as exc:
        logger.debug("[YouTube] oEmbed title fetch failed: %s", exc)
    return f"YouTube video ({video_id})"


def _detect_transcript_language(text: str) -> str:
    """
    Detect transcript language. Returns an ISO 639-1 code or 'unknown'.

    Strategy:
      1. Try langdetect library if installed.
      2. Fallback: fraction of ASCII chars — non-Latin scripts score very low.
    """
    sample = text[:3000]
    try:
        from langdetect import detect  # type: ignore
        lang = detect(sample)
        return lang or "unknown"
    except ImportError:
        pass
    except Exception:
        pass

    # ASCII-fraction heuristic: Urdu/Arabic/Chinese/Devanagari have very few ASCII chars
    if not sample:
        return "unknown"
    ascii_count = sum(1 for c in sample if ord(c) < 128)
    ascii_frac = ascii_count / len(sample)
    if ascii_frac > 0.80:
        return "en"
    # Rough detection: Urdu/Arabic block is U+0600–U+06FF
    arabic_count = sum(1 for c in sample if "؀" <= c <= "ۿ")
    if arabic_count / len(sample) > 0.05:
        return "ur"  # Urdu / Arabic script
    return "unknown"


def _load_youtube(url: str) -> list[Document]:
    """
    Load a YouTube video transcript and return the full text as a single
    LangChain Document ready for semantic chunking.

    Strategy
    --------
    1. Try YouTubeTranscriptApi (>= 0.6 API — uses .fetch() on a
       FetchedTranscript object, not the old list_transcripts() pattern).
    2. Fall back to Whisper audio transcription if the transcript API fails
       (e.g., transcripts disabled for the video).

    Each chunk carries metadata:
      source       : original YouTube URL
      video_id     : 11-character YouTube video ID
      video_title  : human-readable title from YouTube oEmbed API
      language     : detected transcript language (ISO 639-1)
      chunk_index  : 0-based chunk number
      source_type  : "youtube_transcript" or "whisper"
      doc_type     : "youtube"
    """
    video_id = _extract_video_id(url)
    if not video_id:
        raise ValueError(f"Could not extract video ID from URL: {url}")

    full_text: Optional[str] = None
    source_type = "youtube_transcript"

    # ── Attempt 1: YouTubeTranscriptApi ───────────────────────────────────────
    try:
        from youtube_transcript_api import YouTubeTranscriptApi

        try:
            # Modern API (>= 0.6) uses instance method api.fetch(video_id)
            api = YouTubeTranscriptApi()
            fetched = api.fetch(video_id)
            snippets = list(fetched)
            full_text = " ".join(
                s.text if hasattr(s, "text") else s.get("text", "")
                for s in snippets
            ).strip()
        except (TypeError, AttributeError, ValueError) as sub_exc:
            logger.info("[YouTube] Modern fetch API failed (%s) — falling back to get_transcript", sub_exc)
            # Older API (< 0.6) or fallback to classmethod get_transcript
            try:
                raw = YouTubeTranscriptApi.get_transcript(video_id)
                full_text = " ".join(s.get("text", "") for s in raw).strip()
            except Exception:
                raise sub_exc

        if not full_text:
            raise ValueError("Empty transcript returned by YouTubeTranscriptApi")

        logger.info(
            "[YouTube] Transcript loaded for %s (%d chars)", video_id, len(full_text)
        )

    except Exception as yt_exc:
        logger.warning("[YouTube] Transcript API failed: %s — trying Whisper", yt_exc)

        # ── Attempt 2: Whisper audio transcription ────────────────────────────
        whisper_model = _get_whisper_model()
        if whisper_model is None:
            raise RuntimeError(
                f"YouTube transcript unavailable for {url} and Whisper is not installed. "
                "Install 'openai-whisper' or enable video transcripts."
            ) from yt_exc

        try:
            import yt_dlp
        except ImportError:
            raise RuntimeError(
                "yt-dlp is required for Whisper audio download. "
                "Install it with: pip install yt-dlp"
            ) from yt_exc

        with tempfile.TemporaryDirectory(dir=_SAFE_TMP_ROOT) as tmp_dir:
            # Use %(ext)s so yt-dlp fills in the actual extension
            # instead of hardcoding .mp3 which may not exist before conversion
            outtmpl = os.path.join(tmp_dir, "audio.%(ext)s")
            ydl_opts = {
                "format": "bestaudio[ext=m4a]/bestaudio[ext=webm]/bestaudio/best",
                "outtmpl": outtmpl,
                "quiet": True,
                "no_warnings": True,
                # No FFmpeg postprocessor — Whisper reads m4a/webm natively
            }
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=True)
                ext = info.get("ext", "m4a")

            # Find the actual downloaded file
            audio_path = os.path.join(tmp_dir, f"audio.{ext}")
            if not os.path.exists(audio_path):
                # fallback: grab whatever was downloaded
                import glob
                candidates = glob.glob(os.path.join(tmp_dir, "audio.*"))
                if not candidates:
                    raise RuntimeError("yt-dlp downloaded nothing — check the URL or your network.")
                audio_path = candidates[0]

            logger.info("[Whisper] Transcribing %s", audio_path)
            result = whisper_model.transcribe(audio_path)
            full_text = result.get("text", "").strip()
            source_type = "whisper"
            logger.info(
                "[Whisper] Transcription complete for %s (%d chars)", video_id, len(full_text)
            )

    if not full_text.strip():
        raise ValueError(f"No transcript content found for: {url}")

    # Clean filler words and duplicate sentences from auto-generated captions
    original_len = len(full_text)
    full_text = _clean_youtube_transcript(full_text)
    logger.info(
        "[YouTube] Transcript cleaned: %d → %d chars (removed %.0f%%)",
        original_len, len(full_text),
        100 * (1 - len(full_text) / max(original_len, 1)),
    )

    # Detect transcript language — warn if non-English so the user knows
    # that English queries may not retrieve content from this video.
    detected_lang = _detect_transcript_language(full_text)
    if detected_lang not in ("en", "unknown"):
        logger.warning(
            "[YouTube] Non-English transcript detected (lang=%s) for %s. "
            "English queries will likely not match this content.",
            detected_lang, url,
        )
    else:
        logger.info("[YouTube] Transcript language: %s", detected_lang)

    # Fetch human-readable video title
    video_title = _get_video_title(url, video_id)

    # Return the ENTIRE transcript as a single Document.
    # The SemanticChunker in rag.py / api.py will do all chunking — it uses
    # word-window pseudo-sentences for YouTube content (no punctuation) and
    # finds topic-shift boundaries across the full text.  Pre-splitting here
    # with fixed word windows prevents the semantic chunker from detecting any
    # boundary that straddles a word-cut.
    word_count = len(full_text.split())
    logger.info(
        "[YouTube] Returning full transcript as single document: %d words, %d chars, title=%r",
        word_count, len(full_text), video_title,
    )

    return [Document(
        page_content=full_text,
        metadata={
            "source":       url,
            "video_id":     video_id,
            "video_title":  video_title,
            "language":     detected_lang,
            "chunk_index":  0,
            "source_type":  source_type,
            "doc_type":     "youtube",
            "section":      "general",
            "policy_type":  "general",
        },
    )]


# ══════════════════════════════════════════════════════════════════════════════
# PDF LOADER
# ══════════════════════════════════════════════════════════════════════════════

def _load_pdf(file_path: str, filename: str = "") -> list[Document]:
    """
    Load a PDF with per-page Documents.

    Priority order (uses whatever is installed):
      1. pypdf  — modern, pure-Python, actively maintained (pip install pypdf)
      2. pdfplumber — good for complex layouts (pip install pdfplumber)

    Raises RuntimeError if neither library is available.
    """
    # ── pypdf (preferred — installed as 'pypdf' v6+) ──────────────────────────
    try:
        from pypdf import PdfReader  # type: ignore

        docs: list[Document] = []
        reader = PdfReader(file_path)
        total = len(reader.pages)
        for page_num, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": filename or file_path,
                        "filename": filename or os.path.basename(file_path),
                        "page": page_num + 1,
                        "total_pages": total,
                    },
                ))
        logger.info("[PDF/pypdf] Loaded %d pages from '%s'", len(docs), filename)
        return docs

    except ImportError:
        logger.warning("[PDF] pypdf not available, trying pdfplumber")

    # ── pdfplumber (fallback) ─────────────────────────────────────────────────
    try:
        import pdfplumber  # type: ignore

        docs = []
        with pdfplumber.open(file_path) as pdf:
            total = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages):
                text = (page.extract_text() or "").strip()
                if text:
                    docs.append(Document(
                        page_content=text,
                        metadata={
                            "source": filename or file_path,
                            "filename": filename or os.path.basename(file_path),
                            "page": page_num + 1,
                            "total_pages": total,
                        },
                    ))
        logger.info("[PDF/pdfplumber] Loaded %d pages from '%s'", len(docs), filename)
        return docs

    except ImportError:
        pass

    raise RuntimeError(
        "No PDF library installed. Run: pip install pypdf"
    )


# ══════════════════════════════════════════════════════════════════════════════
# DOCX / TXT / CSV / XLSX LOADERS
# ══════════════════════════════════════════════════════════════════════════════

def _load_docx(file_path: str, filename: str = "") -> list[Document]:
    """Load a DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError:
        raise RuntimeError("python-docx is required: pip install python-docx")

    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    if not text:
        return []
    return [Document(
        page_content=text,
        metadata={
            "source": filename or file_path,
            "filename": filename or os.path.basename(file_path),
            "page": 1,
        },
    )]


def _load_txt(file_path: str, filename: str = "") -> list[Document]:
    """Load a plain text file."""
    try:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            text = f.read().strip()
    except Exception as exc:
        raise RuntimeError(f"Could not read text file '{filename}': {exc}") from exc

    if not text:
        return []
    return [Document(
        page_content=text,
        metadata={
            "source": filename or file_path,
            "filename": filename or os.path.basename(file_path),
            "page": 1,
        },
    )]


def _load_csv(file_path: str, filename: str = "") -> list[Document]:
    """Load a CSV as a plain-text representation."""
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
    except Exception as exc:
        raise RuntimeError(f"Could not parse CSV '{filename}': {exc}") from exc

    text = df.to_string(index=False)
    if not text.strip():
        return []
    return [Document(
        page_content=text,
        metadata={
            "source": filename or file_path,
            "filename": filename or os.path.basename(file_path),
            "page": 1,
            "rows": len(df),
            "columns": list(df.columns),
        },
    )]


def _load_xlsx(file_path: str, filename: str = "") -> list[Document]:
    """Load an XLSX file; each sheet becomes one Document."""
    try:
        import pandas as pd
        sheets = pd.read_excel(file_path, sheet_name=None)
    except Exception as exc:
        raise RuntimeError(f"Could not parse Excel file '{filename}': {exc}") from exc

    docs: list[Document] = []
    for sheet_name, df in sheets.items():
        text = df.to_string(index=False).strip()
        if text:
            docs.append(Document(
                page_content=f"[Sheet: {sheet_name}]\n{text}",
                metadata={
                    "source": filename or file_path,
                    "filename": filename or os.path.basename(file_path),
                    "sheet": sheet_name,
                    "page": 1,
                },
            ))
    return docs


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE LOADER (OCR)
# ══════════════════════════════════════════════════════════════════════════════

def _load_image(file_path: str, filename: str = "") -> list[Document]:
    """OCR an image file using Tesseract via pytesseract."""
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except ImportError:
        raise RuntimeError(
            "Tesseract OCR requires Pillow + pytesseract: "
            "pip install Pillow pytesseract"
        )

    image = Image.open(file_path)
    text = pytesseract.image_to_string(image).strip()
    if not text:
        return []
    return [Document(
        page_content=text,
        metadata={
            "source": filename or file_path,
            "filename": filename or os.path.basename(file_path),
            "page": 1,
            "source_type": "ocr",
        },
    )]


# ══════════════════════════════════════════════════════════════════════════════
# AUDIO / VIDEO LOADER (Whisper)
# ══════════════════════════════════════════════════════════════════════════════

def _load_audio(file_path: str, filename: str = "") -> list[Document]:
    """Transcribe audio/video using local Whisper."""
    model = _get_whisper_model()
    if model is None:
        raise RuntimeError(
            "Whisper is required for audio files: pip install openai-whisper"
        )

    result = model.transcribe(file_path)
    text = result.get("text", "").strip()
    if not text:
        return []
    return [Document(
        page_content=text,
        metadata={
            "source": filename or file_path,
            "filename": filename or os.path.basename(file_path),
            "page": 1,
            "source_type": "whisper",
        },
    )]


# ══════════════════════════════════════════════════════════════════════════════
# WEB PAGE LOADER
# ══════════════════════════════════════════════════════════════════════════════

def _parse_html_to_text(html: str) -> tuple[str, str]:
    """
    Parse raw HTML into clean plain text and extract the page title.

    Removes boilerplate tags, prefers <main>/<article>, filters noise lines,
    and deduplicates consecutive identical lines.

    Returns (clean_text, title).
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")

    # ── Remove boilerplate tags ───────────────────────────────────────────────
    for tag in soup(["script", "style", "nav", "footer", "header",
                     "aside", "form", "noscript", "svg", "iframe",
                     "button", "figure", "picture"]):
        tag.decompose()

    # ── Prefer semantic content containers ───────────────────────────────────
    container = (
        soup.find("main")
        or soup.find("article")
        or soup.find("section")
        or soup.find("body")
        or soup
    )
    raw_text = container.get_text(separator="\n", strip=True)

    # ── Apply noise filters ───────────────────────────────────────────────────
    clean_text = _clean_webpage_text(raw_text)

    # ── Extract title ─────────────────────────────────────────────────────────
    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()

    return clean_text, title


def _fetch_with_requests(
    url: str,
    timeout: int,
    user_agent: str,
) -> tuple[str, int]:
    """
    Fetch a URL with requests and return (html_text, status_code).
    Raises RuntimeError with a user-friendly message on HTTP errors.
    """
    import requests

    headers = {"User-Agent": user_agent}
    try:
        resp = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
    except requests.exceptions.ConnectionError as exc:
        raise RuntimeError(
            f"Could not connect to '{url}'. Check the URL and your internet connection."
        ) from exc
    except requests.exceptions.Timeout as exc:
        raise RuntimeError(
            f"Request to '{url}' timed out after {timeout}s. The site may be slow or unreachable."
        ) from exc
    except Exception as exc:
        raise RuntimeError(f"Failed to fetch URL '{url}': {exc}") from exc

    # ── Friendly HTTP error messages ──────────────────────────────────────────
    if resp.status_code == 403:
        raise RuntimeError(
            f"'{url}' blocked the request (HTTP 403 Forbidden). "
            "The site requires login or only allows direct browser access."
        )
    if resp.status_code == 429:
        raise RuntimeError(
            f"'{url}' returned HTTP 429 Too Many Requests. "
            "The site is rate-limiting. Try again in a few minutes."
        )
    if resp.status_code == 404:
        raise RuntimeError(
            f"Page not found (HTTP 404): '{url}'. Check that the URL is correct."
        )
    if resp.status_code >= 500:
        raise RuntimeError(
            f"'{url}' returned a server error (HTTP {resp.status_code}). "
            "The site may be temporarily down."
        )
    if not resp.ok:
        raise RuntimeError(
            f"'{url}' returned HTTP {resp.status_code}."
        )

    # ── Encoding detection ────────────────────────────────────────────────────
    # requests guesses encoding from headers; chardet is more reliable for
    # non-UTF-8 pages (some insurers still serve ISO-8859-1 / Windows-1252).
    detected_enc = _detect_encoding(resp.content)
    try:
        html = resp.content.decode(detected_enc, errors="replace")
    except (LookupError, UnicodeDecodeError):
        html = resp.text  # fallback to requests' own decoding

    return html, resp.status_code


def _fetch_with_playwright(url: str, timeout: int) -> str:
    """
    Fetch a JS-rendered page using Playwright (Chromium headless).

    Playwright is an OPTIONAL dependency — the import is deferred to runtime
    so the rest of the module loads fine even when it is not installed.
    Pylance / pyright "import could not be resolved" warnings are expected and
    harmless; the try/except ImportError handles the missing-package case.

    Install with:
        pip install playwright
        playwright install chromium

    Returns raw HTML string.
    Raises RuntimeError if Playwright is not installed or the page fails to load.
    """
    # Deferred import — optional dependency.  # noqa: PLC0415
    # pyright: ignore[reportMissingModuleSource]
    try:
        import importlib
        _pw_module   = importlib.import_module("playwright.sync_api")  # type: ignore[import-untyped]
        sync_playwright = _pw_module.sync_playwright                   # type: ignore[attr-defined]
        PWTimeout       = _pw_module.TimeoutError                      # type: ignore[attr-defined]
    except ImportError:
        raise RuntimeError(
            "Playwright is not installed. Run: pip install playwright && playwright install chromium"
        )

    logger.info("[Playwright] Fetching JS-rendered page: %s", url)
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            context = browser.new_context(
                user_agent=(
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/120.0.0.0 Safari/537.36"
                )
            )
            page = context.new_page()
            try:
                page.goto(url, wait_until="networkidle", timeout=timeout * 1000)
            except PWTimeout:
                # networkidle can time-out on chatty pages; domcontentloaded is enough
                logger.warning("[Playwright] networkidle timed out — retrying with domcontentloaded")
                page.goto(url, wait_until="domcontentloaded", timeout=timeout * 1000)
            html = page.content()
            browser.close()
        return html
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(f"Playwright failed to load '{url}': {exc}") from exc


_BROWSER_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/125.0.0.0 Safari/537.36"
)


def _load_web_page(
    url: str,
    *,
    timeout: int = 20,
    max_chars: int = 100_000,
    user_agent: str = _BROWSER_UA,
) -> list[Document]:
    """
    Fetch a web page and convert its main text content to a Document.

    Strategy
    --------
    1. Fetch with requests + decode with chardet.
    2. Parse HTML → clean text with BeautifulSoup + noise filters.
    3. If the resulting text is too sparse (< _SPARSE_CONTENT_THRESHOLD chars),
       fall back to Playwright headless browser to render JS-heavy pages.
    4. Apply _clean_webpage_text() to remove nav noise and duplicate lines.
    5. Truncate at max_chars and return as a single Document.
    """
    try:
        import requests  # noqa: F401
        from bs4 import BeautifulSoup  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "requests and beautifulsoup4 are required for web scraping: "
            "pip install requests beautifulsoup4"
        )

    html = None
    title = url

    # ── Step 1: Try requests ──────────────────────────────────────────────────
    try:
        html, _ = _fetch_with_requests(url, timeout=timeout, user_agent=user_agent)
        clean_text, title = _parse_html_to_text(html)
        logger.info(
            "[WebPage/requests] %s → %d chars after cleaning", url, len(clean_text)
        )
    except RuntimeError:
        raise  # re-raise friendly HTTP errors immediately (403, 404, 429, 5xx)
    except Exception as exc:
        logger.warning("[WebPage/requests] Failed for %s: %s — trying Playwright", url, exc)
        clean_text = ""

    # ── Step 2: Playwright fallback if content is too sparse ──────────────────
    if len(clean_text) < _SPARSE_CONTENT_THRESHOLD:
        logger.info(
            "[WebPage] Sparse content (%d chars) from requests — trying Playwright for %s",
            len(clean_text), url,
        )
        try:
            playwright_html = _fetch_with_playwright(url, timeout=max(timeout, 30))
            pw_text, pw_title = _parse_html_to_text(playwright_html)
            if len(pw_text) > len(clean_text):
                clean_text = pw_text
                title = pw_title or title
                logger.info(
                    "[WebPage/Playwright] %s → %d chars after cleaning", url, len(clean_text)
                )
            else:
                logger.warning(
                    "[WebPage/Playwright] No improvement for %s (%d chars)", url, len(pw_text)
                )
        except RuntimeError as pw_exc:
            logger.warning("[WebPage/Playwright] Skipped: %s", pw_exc)

    if not clean_text:
        return []

    # ── Step 3: Truncate and build Document ───────────────────────────────────
    clean_text = clean_text[:max_chars]

    return [Document(
        page_content=clean_text,
        metadata={
            "source": url,
            "source_url": url,
            "filename": url,
            "title": title,
            "page": 1,
            "source_type": "web",
            "doc_type": "general",
            "char_count": len(clean_text),
        },
    )]


# ══════════════════════════════════════════════════════════════════════════════
# MAIN ENTRY POINTS
# ══════════════════════════════════════════════════════════════════════════════

def load_document(file_path: str, filename: str = "") -> list[Document]:
    """
    Load a local file and return a list of LangChain Documents.

    Dispatches to the appropriate loader based on file extension.
    Raises FileValidationError for unsupported or oversized files.
    Raises RuntimeError when a required library is missing.

    Parameters
    ----------
    file_path : Absolute path to the local file (e.g. a temp file).
    filename  : Original filename (used for metadata and extension detection).
                Defaults to basename of file_path.
    """
    fname = filename or os.path.basename(file_path)
    ext = Path(fname).suffix.lower()

    # ── Size check ────────────────────────────────────────────────────────────
    try:
        size = os.path.getsize(file_path)
        if size > MAX_FILE_SIZE_BYTES:
            raise FileValidationError(
                f"File '{fname}' is {size // (1024*1024)} MB — exceeds the "
                f"{MAX_FILE_SIZE_BYTES // (1024*1024)} MB limit."
            )
    except FileNotFoundError:
        raise FileValidationError(f"File not found: {file_path}")

    # ── Extension check ───────────────────────────────────────────────────────
    if ext not in ALLOWED_EXTENSIONS:
        raise FileValidationError(
            f"Unsupported file type '{ext}'. Allowed: {sorted(ALLOWED_EXTENSIONS)}"
        )

    # ── Dispatch to loader ────────────────────────────────────────────────────
    if ext == ".pdf":
        return _load_pdf(file_path, fname)
    if ext in (".docx", ".doc"):
        return _load_docx(file_path, fname)
    if ext == ".txt":
        return _load_txt(file_path, fname)
    if ext == ".csv":
        return _load_csv(file_path, fname)
    if ext in (".xlsx", ".xls"):
        return _load_xlsx(file_path, fname)
    if ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"):
        return _load_image(file_path, fname)
    if ext in (".mp3", ".mp4", ".wav", ".m4a", ".ogg", ".webm"):
        return _load_audio(file_path, fname)

    raise FileValidationError(f"No loader implemented for extension '{ext}'")


def load_url(url: str) -> list[Document]:
    """
    Load content from a URL.

    Routing:
      - YouTube URL → _load_youtube()
      - Everything else → fetch HTML, strip tags, return as Document

    Returns a list of Documents (multiple for YouTube chunks, usually one
    for web pages).
    """
    if is_youtube_url(url):
        return _load_youtube(url)
    return _load_web_page(url)


def load_url_advanced(
    url: str,
    *,
    timeout: int = 30,
    max_chars: int = 100_000,
    user_agent: str = _BROWSER_UA,
) -> list[Document]:
    """
    Advanced URL loader with configurable timeout, size limit, and user-agent.

    Identical routing to load_url() (YouTube vs web page) but exposes
    additional parameters for callers that need finer control.
    """
    if is_youtube_url(url):
        return _load_youtube(url)
    return _load_web_page(url, timeout=timeout, max_chars=max_chars, user_agent=user_agent)